"""
Jarvis Web Server (updated)
---------------------------
Naye features:
- File/Image/PDF upload support in chat
- Search context Jarvis ko milta hai
- brain.py/memory.py/tools.py seedha import
"""

from flask import Flask, render_template, request, jsonify, send_file
import base64
import os
import tempfile
import shutil
import subprocess
import sys
import re
import json
import memory
import brain
import rag
import voice
import phone_bridge
import twilio_call  # Twilio phone-call assistant (/voice, /respond)
import keepalive    # Render free-tier ko sleep hone se rokne wala self-ping
import memory_guard # Render free-tier OOM-crash se bachaane wala watchdog
import scheduler    # APScheduler — background/recurring maintenance jobs
from logger import get_logger
log = get_logger("server")

app = Flask(__name__)
HISTORY_LIMIT = 16

# Ek single upload/request 25MB se zyada na ho — warna bada file/image
# upload akele hi Render ke 512MB RAM limit ko cross kar sakta hai aur
# poore worker ko crash kar sakta hai. Flask khud hi 413 error de dega.
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024  # 25MB

# Twilio call routes register karo (isi Flask app par — Render par ek hi
# service se web-chat aur phone-call dono chalte hain).
twilio_call.register_routes(app)

# Keep-alive thread module-level par start karo — gunicorn 'server:app' ko
# seedha import karta hai, __main__ guard kabhi nahi chalta, isliye yahan
# top-level par hi start karna zaroori hai.
keepalive.start_keepalive_thread()

# Memory watchdog bhi yahin top-level par start karo — Render free plan ka
# 512MB RAM limit cross hone se PEHLE hi proactively gc.collect()/graceful
# self-restart kar deta hai, taaki Render ka apna abrupt OOM-kill kabhi
# trigger hi na ho (dekho memory_guard.py docstring).
memory_guard.start_memory_guard()

# APScheduler bhi yahin top-level par start karo (same reason — gunicorn
# seedha 'server:app' import karta hai, __main__ guard nahi chalta).
scheduler.register_default_jobs()


# Koi bhi route mein agar kabhi ek anhandled exception bach jaaye, to
# poora worker crash hone ke bajaye ek clean JSON 500 milta hai aur log
# ho jaata hai — Jarvis chalta rehta hai, sirf woh ek request fail hoti
# hai. Normal HTTP errors (404/405/etc, abort() calls) apne asli status
# code ke saath hi pass hote hain — sirf GENUINE unexpected crashes yahan
# 500 mein convert hote hain.
from werkzeug.exceptions import HTTPException

@app.errorhandler(Exception)
def _handle_uncaught_error(e):
    if isinstance(e, HTTPException):
        return e
    log.exception(f"Uncaught error on {request.path}: {e}")
    return jsonify({"error": "Kuch galat ho gaya, ek baar phir try karo."}), 500


@app.errorhandler(413)
def _handle_too_large(e):
    return jsonify({"error": "File/request bahut badi hai (max 25MB)."}), 413
scheduler.start()


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/api/ping")
def ping():
    """
    Lightweight health-check endpoint — koi auth nahi lagti, kuch nahi karta,
    bas 'alive' bol deta hai. Isi ko keepalive.py background thread har
    10 minute mein khud call karta hai taaki Render free-tier service kabhi
    15-minute-inactivity ki wajah se sleep na ho, aur jab bhi call aaye,
    AI turant available rahe.
    """
    import time as _time_ping
    return jsonify({"status": "alive", "ts": _time_ping.time()})


@app.route("/service-worker.js")
def service_worker():
    """
    Service worker root se serve karta hai — kuch browsers strict scope
    check karte hain, isliye /static/ ke alawa /service-worker.js bhi rakha.
    """
    from flask import send_from_directory, Response
    resp = send_from_directory(
        os.path.join(os.path.dirname(__file__), "static"), "service-worker.js")
    resp.headers["Service-Worker-Allowed"] = "/"
    resp.headers["Content-Type"] = "application/javascript"
    return resp


# ---------- Chat list ----------

@app.route("/api/chats", methods=["GET"])
def get_chats():
    return jsonify({"chats": memory.list_chats()})

@app.route("/api/chats/new", methods=["POST"])
def new_chat():
    chat_id = memory.create_chat()
    return jsonify({"chat_id": chat_id})

@app.route("/api/chats/<chat_id>", methods=["GET"])
def get_chat(chat_id):
    messages = memory.load_chat(chat_id)
    return jsonify({"messages": messages})

@app.route("/api/chats/<chat_id>", methods=["DELETE"])
def delete_chat_route(chat_id):
    memory.delete_chat(chat_id)
    return jsonify({"success": True})


# ---------- Message (with optional file attachments) ----------

@app.route("/api/generate_bg_image", methods=["POST"])
def generate_bg_image():
    """
    AI se background image generate karta hai — chat history mein save NAHI
    hota, user ko chat bubble mein nahi dikhta. Sirf image path return karta hai.
    """
    data = request.get_json() or {}
    desc = (data.get("description") or "").strip()
    if not desc:
        return jsonify({"success": False, "message": "Description do."})

    prompt = (
        f'Generate an image for a chat app background. Description: "{desc}". '
        f'Make it aesthetic, not too bright, suitable as a background.'
    )

    try:
        response = brain.ask_jarvis([], prompt, chat_id="__bg_gen__")
        import re as _re4
        match = _re4.search(r'IMAGE_GENERATED:(/static/generated/[a-zA-Z0-9_.\-]+)', response)
        if not match:
            return jsonify({"success": False, "message": "Image generate nahi hui, dobara try karo."})
        return jsonify({"success": True, "path": match.group(1)})
    except Exception as e:
        return jsonify({"success": False, "message": f"Error: {e}"})


@app.route("/api/generate_theme", methods=["POST"])
def generate_theme():
    """
    AI se theme colors generate karta hai — chat history mein save NAHI hota,
    user ko chat bubble mein bhi nahi dikhta. Sirf JSON colors return karta hai.
    """
    data = request.get_json() or {}
    desc = (data.get("description") or "").strip()
    if not desc:
        return jsonify({"success": False, "message": "Description do."})

    prompt = (
        f'Generate a JSON theme object for a dark chat UI. Description: "{desc}". '
        f'Return ONLY valid JSON (no markdown, no explanation, no code fences) with exactly these keys: '
        f'{{"--bg","--surface","--line","--text","--text-dim","--accent","--accent2","--msg-user","--msg-jarvis"}}. '
        f'All values must be valid CSS hex colors. Make it look cool and match the description. '
        f'Respond with ONLY the JSON object, nothing else.'
    )

    try:
        # Chat history khaali bhejo — yeh isolated request hai, conversation thread nahi
        response = brain.ask_jarvis([], prompt, chat_id="__theme_gen__")

        # JSON extract karo response se
        import re as _re3
        json_match = _re3.search(r'\{[\s\S]*?\}', response)
        if not json_match:
            return jsonify({"success": False, "message": "AI se valid theme nahi mila, dobara try karo."})

        theme_vars = json.loads(json_match.group(0))
        required = ['--bg','--surface','--line','--text','--text-dim','--accent','--accent2','--msg-user','--msg-jarvis']
        if not all(k in theme_vars for k in required):
            return jsonify({"success": False, "message": "Theme incomplete hai, dobara try karo."})

        return jsonify({"success": True, "vars": theme_vars})
    except json.JSONDecodeError:
        return jsonify({"success": False, "message": "AI ne galat format diya, dobara try karo."})
    except Exception as e:
        return jsonify({"success": False, "message": f"Error: {e}"})


@app.route("/api/chats/<chat_id>/message", methods=["POST"])
def send_message(chat_id):
    data = request.get_json()
    user_text = (data or {}).get("message", "").strip()
    files = (data or {}).get("files", [])  # [{name, type, base64}, ...]

    if not user_text and not files:
        return jsonify({"reply": "Kuch toh likho ya file bhejo!"})

    messages = memory.load_chat(chat_id)

    # Bootstrap key commands
    bootstrap_key_cmd = memory.try_extract_api_command(user_text)
    bootstrap_names = set(memory.GROQ_KEY_NAMES) | set(memory.GEMINI_KEY_NAMES) | set(memory.OPENROUTER_KEY_NAMES)
    stream_handled, stream_response = (False, None) if files else brain.handle_stream_command(user_text)

    if bootstrap_key_cmd and bootstrap_key_cmd[0] in bootstrap_names:
        memory.save_secret(bootstrap_key_cmd[0], bootstrap_key_cmd[1])
        response = f"Theek hai, maine '{bootstrap_key_cmd[0]}' API key yaad rakh li hai."
    elif stream_handled:
        # play/status/stop stream commands — deterministic, LLM round-trip
        # ki zarurat nahi. Non-blocking hai (Popen-based), isliye request
        # turant return hoti hai, video khatam hone tak wait nahi karti.
        response = stream_response
    else:
        # File attachments process karo — sab types support
        combined_text = user_text or ""
        if files:
            file_context = _process_files(files)
            if not combined_text:
                combined_text = "Yeh file/attachment dekho aur seedha batao isme kya hai."
            combined_text = file_context + "\n\n[User ka sawaal]: " + combined_text

        # chat_id pass karo brain ko (RAG ke liye)
        response = brain.ask_jarvis(messages, combined_text, chat_id=chat_id)

    # History mein sirf display text save karo
    if files:
        file_names = ", ".join(f.get("name","file") for f in files)
        display_text = f"{user_text} [{file_names}]" if user_text else f"[{file_names}]"
    else:
        display_text = user_text

    messages.append({"role": "user", "content": display_text})
    messages.append({"role": "assistant", "content": response})
    messages = messages[-HISTORY_LIMIT:]
    memory.save_chat(chat_id, messages)

    return jsonify({"reply": response})


def _process_files(files):
    """
    Files ko process karke Jarvis ke liye context text banata hai.
    Sab file types support karta hai.
    """
    parts = []
    for f in files:
        name = f.get("name", "file")
        ftype = f.get("type", "file")
        b64 = f.get("base64", "")

        if ftype == "image":
            parts.append(f"[USER NE EK IMAGE BHEJI HAI: '{name}' — NEECHE IMAGE HAI, ISKO DHYAN SE DEKHO AUR DESCRIBE KARO]")
            parts.append(f"__IMAGE_ATTACHMENT__:{b64}")
        elif ftype == "pdf":
            text = _extract_pdf_text(b64, name)
            parts.append(f"[USER NE EK PDF BHEJI HAI: '{name}']\nPDF CONTENT:\n{text[:3000]}")
        elif ftype == "text" or ftype == "code" or ftype == "json":
            try:
                if "," in b64:
                    b64_data = b64.split(",", 1)[1]
                else:
                    b64_data = b64
                text = base64.b64decode(b64_data).decode("utf-8", errors="ignore")
                parts.append(f"[USER NE EK TEXT FILE BHEJI HAI: '{name}']\nCONTENT:\n{text[:3000]}")
            except Exception as e:
                parts.append(f"[USER NE FILE BHEJI HAI: '{name}' — text read nahi ho paya: {e}]")
        elif ftype == "doc":
            try:
                if "," in b64:
                    b64_data = b64.split(",", 1)[1]
                else:
                    b64_data = b64
                doc_bytes = base64.b64decode(b64_data)
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    tmp.write(doc_bytes)
                    tmp_path = tmp.name
                try:
                    import subprocess
                    result = subprocess.run(
                        ["docx2txt", tmp_path, "-"],
                        capture_output=True, text=True, timeout=10)
                    if result.returncode == 0 and result.stdout.strip():
                        parts.append(f"[USER NE WORD DOC BHEJI HAI: '{name}']\nCONTENT:\n{result.stdout[:3000]}")
                    else:
                        parts.append(f"[USER NE WORD DOC BHEJI HAI: '{name}' — text extract nahi ho paya]")
                finally:
                    import os
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        log.exception("failed to delete temp docx file")
            except Exception as e:
                parts.append(f"[USER NE DOC FILE BHEJI HAI: '{name}' — error: {e}]")
        elif ftype == "zip":
            # ZIP ka manifest list karo
            try:
                if "," in b64:
                    b64_data = b64.split(",", 1)[1]
                else:
                    b64_data = b64
                zip_bytes = base64.b64decode(b64_data)
                import zipfile, io
                with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zf:
                    names_list = zf.namelist()
                    # Choti text files read karo
                    readable = []
                    for zname in names_list[:20]:
                        if zname.endswith(("/", "\\")):
                            continue
                        ext = zname.lower().rsplit(".", 1)[-1]
                        if ext in ("py","js","ts","html","css","json","txt","md","yaml","yml","sh","env","cfg","ini","xml","csv"):
                            try:
                                txt = zf.read(zname).decode("utf-8", errors="ignore")
                                readable.append(f"--- {zname} ---\n{txt[:800]}")
                            except Exception:
                                log.exception("unexpected error - see memory/jarvis_errors.log")
                    summary = f"Files ({len(names_list)}): " + ", ".join(names_list[:30])
                    content_str = "\n\n".join(readable[:5]) if readable else "(binary files only)"
                    parts.append(f"[USER NE ZIP FILE BHEJI HAI: '{name}']\n{summary}\n\nReadable content:\n{content_str}")
            except Exception as e:
                parts.append(f"[USER NE ZIP FILE BHEJI HAI: '{name}' — error: {e}]")
        elif ftype == "video":
            parts.append(f"[USER NE EK VIDEO FILE BHEJI HAI: '{name}' — video play karne ke liye download karo]")
        elif ftype == "audio":
            parts.append(f"[USER NE AUDIO FILE BHEJI HAI: '{name}']")
        elif ftype == "spreadsheet":
            try:
                if "," in b64:
                    b64_data = b64.split(",", 1)[1]
                else:
                    b64_data = b64
                sheet_bytes = base64.b64decode(b64_data)
                if name.lower().endswith(".csv"):
                    text = sheet_bytes.decode("utf-8", errors="ignore")
                    lines_csv = text.strip().split("\n")[:30]
                    parts.append(f"[USER NE CSV FILE BHEJI HAI: '{name}']\nContent (pehle 30 rows):\n" + "\n".join(lines_csv))
                else:
                    try:
                        import openpyxl, io
                        wb = openpyxl.load_workbook(io.BytesIO(sheet_bytes), read_only=True)
                        rows_text = []
                        ws = wb.active
                        for i, row in enumerate(ws.iter_rows(values_only=True)):
                            if i > 25: break
                            rows_text.append("\t".join(str(c or "") for c in row))
                        parts.append(f"[USER NE EXCEL FILE BHEJI HAI: '{name}']\nContent:\n" + "\n".join(rows_text))
                    except Exception:
                        parts.append(f"[USER NE SPREADSHEET BHEJI HAI: '{name}' — Excel file, openpyxl se read karo]")
            except Exception as e:
                parts.append(f"[USER NE SPREADSHEET BHEJI HAI: '{name}' — error: {e}]")
        else:
            parts.append(f"[USER NE EK FILE BHEJI HAI: '{name}' (type: {ftype})]")

    return "\n\n".join(parts)


def _extract_pdf_text(base64_data, filename):
    """PDF se text nikalta hai."""
    try:
        # base64 decode
        if "," in base64_data:
            base64_data = base64_data.split(",", 1)[1]
        pdf_bytes = base64.b64decode(base64_data)

        # Temp file mein save karo
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name

        try:
            import subprocess
            result = subprocess.run(
                ["pdftotext", tmp_path, "-"],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            log.exception("unexpected error - see memory/jarvis_errors.log")

        # PyPDF2 fallback
        try:
            import PyPDF2
            import io
            reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            text = ""
            for page in reader.pages[:10]:  # Max 10 pages
                text += page.extract_text() or ""
            if text.strip():
                return text.strip()
        except Exception:
            log.exception("unexpected error - see memory/jarvis_errors.log")

        return f"(PDF '{filename}' ka text extract nahi ho paya, lekin file receive hui.)"
    except Exception as e:
        return f"(PDF process karne mein error: {e})"


# ---------- API keys ----------

@app.route("/api/save_key", methods=["POST"])
def save_key():
    data = request.get_json()
    name = (data or {}).get("name", "").strip()
    value = (data or {}).get("value", "").strip()
    if not name or not value:
        return jsonify({"success": False, "message": "Naam aur value dono chahiye."})
    standard_name = memory.normalize_api_name(name)
    memory.save_secret(standard_name, value)
    return jsonify({"success": True, "message": f"'{standard_name}' key save ho gayi."})

@app.route("/api/keys", methods=["GET"])
def list_keys():
    return jsonify({"keys": memory.list_known_secrets()})


# ---------- Health / Status Dashboard ----------

@app.route("/api/status", methods=["GET"])
def status_dashboard():
    """
    Ek jagah pe pura Jarvis health check — kaunsi API keys available hain
    (count only, values kabhi nahi), RAG memory stats, active persona,
    aur selected model. Koi live API call nahi karta (quota bachane ke
    liye) — sirf jo memory mein already stored hai woh dikhata hai.
    """
    try:
        groq_keys   = memory.get_available_groq_keys()
        gemini_keys = memory.get_available_gemini_keys()
        or_keys     = memory.get_available_openrouter_keys()

        try:
            rag_stats = rag.get_rag_stats()
        except Exception:
            rag_stats = {"available": False}

        try:
            active_persona = memory.get_active_persona()
        except Exception:
            active_persona = None

        try:
            usage = brain.get_usage_counters()
        except Exception:
            usage = {}

        return jsonify({
            "providers": {
                "groq":       {"keys_available": len(groq_keys)},
                "gemini":     {"keys_available": len(gemini_keys)},
                "openrouter": {"keys_available": len(or_keys)},
            },
            "total_keys": len(groq_keys) + len(gemini_keys) + len(or_keys),
            "selected_model": memory.get_selected_model(),
            "rag": rag_stats,
            "active_persona": active_persona,
            "usage_this_session": usage,
            "ok": (len(groq_keys) + len(gemini_keys) + len(or_keys)) > 0,
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------- Models ----------

@app.route("/api/models", methods=["GET"])
def list_models():
    return jsonify({"models": brain.list_available_models(), "selected": memory.get_selected_model()})

@app.route("/api/models/select", methods=["POST"])
def select_model():
    data = request.get_json()
    model_id = (data or {}).get("model_id", "auto")
    memory.set_selected_model(model_id)
    return jsonify({"success": True, "selected": model_id})


# ---------- Radio Stream Proxy ----------

@app.route("/api/radiostream")
def radio_stream():
    """Radio stream ko proxy karo — CORS bypass ke liye"""
    import urllib.request
    import ssl as _ssl3

    url = request.args.get("url", "").strip()
    if not url or not url.startswith("http"):
        return "Invalid URL", 400

    ctx3 = _ssl3.create_default_context()
    ctx3.check_hostname = False
    ctx3.verify_mode = _ssl3.CERT_NONE

    headers_r = {
        "User-Agent": "Mozilla/5.0",
        "Accept": "*/*",
        "Icy-MetaData": "0",
    }

    from flask import Response, stream_with_context
    try:
        req = urllib.request.Request(url, headers=headers_r)
        opener = urllib.request.build_opener(urllib.request.HTTPSHandler(context=ctx3))
        resp = opener.open(req, timeout=15)
        ctype = resp.headers.get("Content-Type", "audio/mpeg")

        def generate():
            try:
                while True:
                    chunk = resp.read(8192)
                    if not chunk:
                        break
                    yield chunk
            except Exception:
                log.exception("unexpected error - see memory/jarvis_errors.log")

        r = Response(stream_with_context(generate()), content_type=ctype)
        r.headers['Access-Control-Allow-Origin'] = '*'
        r.headers['Cache-Control'] = 'no-cache'
        return r
    except Exception as e:
        from flask import redirect
        return redirect(url, code=302)


# ---------- File Download ----------

import re as _re
import hashlib as _hashlib
import time as _time

# In-memory file store {token: {filename, content, created}}
_file_store = {}

@app.route("/api/create_file", methods=["POST"])
def create_file_route():
    """
    Brain ke response se extracted file ko store karke download token deta hai.
    Plain text formats ke liye seedha text store hota hai.
    Binary formats (pdf, docx, xlsx, zip) ke liye text content se asli
    file generate hoti hai (PDF banta hai PyPDF/reportlab se, etc).
    """
    data = request.get_json()
    filename = (data or {}).get("filename", "file.txt").strip()
    content_text = (data or {}).get("content", "")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    binary_bytes = None
    error = None

    if ext == "pdf":
        binary_bytes, error = _generate_pdf_bytes(content_text)
    elif ext == "docx":
        binary_bytes, error = _generate_docx_bytes(content_text)
    elif ext in ("xlsx",):
        binary_bytes, error = _generate_xlsx_bytes(content_text)
    elif ext == "zip":
        binary_bytes, error = _generate_zip_bytes(content_text)

    token = _hashlib.md5(f"{filename}{_time.time()}".encode()).hexdigest()[:12]

    if binary_bytes is not None:
        _file_store[token] = {"filename": filename, "binary": binary_bytes, "created": _time.time()}
    else:
        _file_store[token] = {"filename": filename, "content": content_text, "created": _time.time()}

    # Purani files clean karo (1 ghante se zyada purani)
    old_tokens = [t for t, v in _file_store.items() if _time.time() - v["created"] > 3600]
    for t in old_tokens:
        del _file_store[t]

    resp = {"token": token, "filename": filename}
    if error:
        resp["warning"] = error
    return jsonify(resp)


def _generate_pdf_bytes(text_content: str):
    """
    Plain text/markdown se ek properly formatted PDF banata hai.
    reportlab use karta hai — Termux mein lightweight install hota hai.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.enums import TA_LEFT
        import io
        import html as _html

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                 topMargin=0.8*inch, bottomMargin=0.8*inch,
                                 leftMargin=0.8*inch, rightMargin=0.8*inch)
        styles = getSampleStyleSheet()
        body_style = styles["Normal"]
        body_style.fontSize = 11
        body_style.leading = 16

        story = []
        for line in text_content.split("\n"):
            line = line.rstrip()
            safe = _html.escape(line) if line else "&nbsp;"
            if line.startswith("# "):
                story.append(Paragraph(_html.escape(line[2:]), styles["Heading1"]))
            elif line.startswith("## "):
                story.append(Paragraph(_html.escape(line[3:]), styles["Heading2"]))
            elif line.startswith("### "):
                story.append(Paragraph(_html.escape(line[4:]), styles["Heading3"]))
            else:
                story.append(Paragraph(safe, body_style))
            story.append(Spacer(1, 4))

        doc.build(story)
        return buf.getvalue(), None
    except ImportError:
        return None, "PDF generate karne ke liye 'reportlab' install karo: pip install reportlab --break-system-packages"
    except Exception as e:
        return None, f"PDF generate error: {e}"


def _generate_docx_bytes(text_content: str):
    """Plain text/markdown se Word document banata hai (python-docx use karke)."""
    try:
        from docx import Document
        import io

        doc = Document()
        for line in text_content.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), None
    except ImportError:
        return None, "DOCX generate karne ke liye 'python-docx' install karo: pip install python-docx --break-system-packages"
    except Exception as e:
        return None, f"DOCX generate error: {e}"


def _generate_xlsx_bytes(text_content: str):
    """
    CSV-style text (comma ya tab separated) se Excel file banata hai.
    Har line ek row hai.
    """
    try:
        import openpyxl
        import io

        wb = openpyxl.Workbook()
        ws = wb.active
        for line in text_content.strip().split("\n"):
            if "\t" in line:
                cells = line.split("\t")
            else:
                cells = line.split(",")
            ws.append(cells)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue(), None
    except Exception as e:
        return None, f"XLSX generate error: {e}"


def _generate_zip_bytes(text_content: str):
    """
    Multi-file ZIP banata hai. Format:
    ===FILE: filename1.py===
    content here...
    ===FILE: filename2.txt===
    more content...
    Agar yeh format nahi mila to poora content ek hi readme.txt mein daal deta hai.
    """
    try:
        import zipfile
        import io
        import re as _re2

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            parts = _re2.split(r"===FILE:\s*([^\n=]+)===\n", text_content)
            if len(parts) > 1:
                # parts[0] is preamble (ignore), then pairs of (filename, content)
                for i in range(1, len(parts), 2):
                    fname = parts[i].strip()
                    fcontent = parts[i+1] if i+1 < len(parts) else ""
                    zf.writestr(fname, fcontent)
            else:
                zf.writestr("readme.txt", text_content)

        return buf.getvalue(), None
    except Exception as e:
        return None, f"ZIP generate error: {e}"


@app.route("/api/download/<token>")
def download_file(token):
    """Token se file download karo — text ya binary dono support."""
    file_info = _file_store.get(token)
    if not file_info:
        return "File nahi mili ya expire ho gayi.", 404
    from flask import Response
    filename = file_info["filename"]

    # Binary file hai (PDF, DOCX, XLSX, ZIP)
    if "binary" in file_info:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
        binary_mime_map = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "zip": "application/zip",
        }
        mime = binary_mime_map.get(ext, "application/octet-stream")
        resp = Response(file_info["binary"], content_type=mime)
        resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp

    # Plain text file
    content_text = file_info["content"]
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    mime_map = {
        "py": "text/x-python", "js": "application/javascript",
        "html": "text/html", "css": "text/css", "json": "application/json",
        "csv": "text/csv", "txt": "text/plain", "md": "text/markdown",
        "sh": "text/x-sh", "xml": "application/xml", "yaml": "text/yaml",
        "yml": "text/yaml", "sql": "text/x-sql", "java": "text/x-java",
        "cpp": "text/x-c++src", "c": "text/x-csrc", "ts": "text/typescript",
    }
    mime = mime_map.get(ext, "application/octet-stream")
    resp = Response(content_text.encode("utf-8"), content_type=mime)
    resp.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


# ---------- TTS ----------

@app.route("/api/tts", methods=["POST"])
def text_to_speech():
    """
    Text-to-speech — subprocess mein isolate kiya hua hai. Termux ke kuch
    builds mein edge-tts/pydub ka native audio decode crash kar sakta hai
    (android ndk-context error) — subprocess isolation se yeh poore server
    ko crash nahi kar paata, sirf TTS request fail hoti hai.

    Persona active ho to uski voice_gender khud-ba-khud use hoti hai (user
    ke manual male/female toggle se bhi zyada priority — taaki roleplay
    character jaisa (boy/girl) hai, awaaz bhi waisi lage). Emotion bhi
    text ke content se khud detect hoke rate/pitch adjust karta hai.
    """
    data = request.get_json()
    text = (data or {}).get("text", "").strip()
    gender = (data or {}).get("gender", "male").strip()
    if not text:
        return jsonify({"error": "Koi text nahi mila"}), 400

    # Active persona ho to uski voice_gender priority leti hai
    try:
        active_persona = memory.get_active_persona()
        if active_persona and active_persona.get("voice_gender"):
            gender = active_persona["voice_gender"]
    except Exception:
        pass

    # Emotion se rate/pitch detect karo (full emotion ke saath bolne jaisa)
    try:
        rate, pitch = voice.detect_emotion_prosody(text)
    except Exception:
        rate, pitch = "+0%", "+0Hz"

    # ── TTS Cache ──
    # Same text + same voice + same emotion pehle bhi bola gaya ho, to
    # dobara generate karne ki jagah cached MP3 seedha bhej do — Render
    # free-tier pe subprocess spin-up slow hota hai, isliye common
    # greetings/phrases ke liye yeh kaafi fast response deta hai.
    import hashlib as _hashlib
    cache_key = _hashlib.sha256(f"{text}|{gender}|{rate}|{pitch}".encode("utf-8")).hexdigest()
    cache_dir = os.path.join(tempfile.gettempdir(), "jarvis_tts_cache")
    os.makedirs(cache_dir, exist_ok=True)
    cache_path = os.path.join(cache_dir, f"{cache_key}.mp3")
    if os.path.isfile(cache_path):
        return send_file(cache_path, mimetype="audio/mpeg")

    try:
        import subprocess as _sp
        import sys as _sys
        import tempfile as _tf

        out_path = os.path.join(_tf.gettempdir(), f"jarvis_tts_subproc_{os.getpid()}.mp3")
        script = (
            "import sys, asyncio, edge_tts\n"
            "text = sys.argv[1]\n"
            "voice = sys.argv[2]\n"
            "out = sys.argv[3]\n"
            "rate = sys.argv[4]\n"
            "pitch = sys.argv[5]\n"
            "async def run():\n"
            "    try:\n"
            "        await edge_tts.Communicate(text, voice, rate=rate, pitch=pitch).save(out)\n"
            "    except Exception:\n"
            "        await edge_tts.Communicate(text, voice).save(out)\n"
            "asyncio.run(run())\n"
        )
        voice_name = "hi-IN-SwaraNeural" if gender == "female" else "hi-IN-MadhurNeural"

        result = _sp.run(
            [_sys.executable, "-c", script, text, voice_name, out_path, rate, pitch],
            capture_output=True, timeout=30, cwd=os.path.dirname(__file__)
        )

        if result.returncode != 0 or not os.path.exists(out_path):
            err_msg = result.stderr.decode("utf-8", errors="ignore")[-300:]
            return jsonify({"error": f"TTS generate nahi hui: {err_msg}"}), 500

        # Cache mein bhi save kar do future requests ke liye
        try:
            shutil.copyfile(out_path, cache_path)
        except Exception:
            pass  # cache fail hone se actual response block nahi hona chahiye

        return send_file(out_path, mimetype="audio/mpeg")
    except _sp.TimeoutExpired:
        return jsonify({"error": "TTS timeout ho gaya."}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/videoproxy")
def video_proxy():
    """
    Pixabay aur doosri sites ki MP4 videos proxy karta hai.
    Range requests support karta hai — video seek karne ke liye zaroori.
    """
    import urllib.request
    import urllib.parse as uparse

    url = request.args.get("url", "").strip()
    if not url or not url.startswith("http"):
        return "Invalid URL", 400

    headers = {
        "User-Agent": "Mozilla/5.0 (Linux; Android 13; Pixel 7 Pro) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36",
        "Accept": "video/webm,video/mp4,video/*;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": f"{uparse.urlparse(url).scheme}://{uparse.urlparse(url).netloc}/",
    }

    # Range request support (video seek ke liye)
    range_header = request.headers.get("Range")
    if range_header:
        headers["Range"] = range_header

    try:
        import ssl as _ssl2
        ctx2 = _ssl2.create_default_context()
        ctx2.check_hostname = False
        ctx2.verify_mode = _ssl2.CERT_NONE
        req = urllib.request.Request(url, headers=headers)
        opener = urllib.request.build_opener(
            urllib.request.HTTPRedirectHandler(),
            urllib.request.HTTPSHandler(context=ctx2)
        )
        with opener.open(req, timeout=30) as resp:
            data = resp.read()
            ctype = resp.headers.get("Content-Type", "video/mp4")
            status = resp.status

        from flask import Response
        r = Response(data, status=status, content_type=ctype)
        r.headers['Accept-Ranges'] = 'bytes'
        r.headers['Access-Control-Allow-Origin'] = '*'
        r.headers['Cache-Control'] = 'public, max-age=3600'
        return r
    except Exception as e:
        # Proxy fail — seedha redirect
        from flask import redirect
        return redirect(url, code=302)


@app.route("/api/hlsproxy")
def hls_proxy():
    """
    HLS (.m3u8) manifests aur unke segments (.ts/.m4s/.aac/keys) ko proxy
    karta hai taaki chat ke andar <video> + hls.js se koi bhi stream CORS
    error ke bina play ho sake — koi bhi external link kholne ki zarurat
    nahi padti.

    Master/media playlists ke andar har URI (variant streams, segments,
    #EXT-X-KEY/#EXT-X-MAP ke URI="..." attributes) is proxy ke through
    hi wapas rewrite ho jaate hain, isliye nested/live playlists aur
    encrypted streams bhi seedha kaam karte hain.
    """
    import re as _re
    import ssl as _ssl3
    import urllib.parse as uparse
    import urllib.request
    from flask import Response, stream_with_context

    url = request.args.get("url", "").strip()
    if not url or not url.lower().startswith(("http://", "https://")):
        return "Invalid URL", 400

    parsed = uparse.urlparse(url)
    headers = {
        "User-Agent": "Mozilla/5.0 (Linux; Android 13; Pixel 7 Pro) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36",
        "Accept": "*/*",
        "Referer": f"{parsed.scheme}://{parsed.netloc}/",
    }
    range_header = request.headers.get("Range")
    if range_header:
        headers["Range"] = range_header

    ctx = _ssl3.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = _ssl3.CERT_NONE

    # MEMORY FIX (Render OOM): manifests chhote text files hain (kuch KB) —
    # poora RAM mein read karna theek hai. Segments (.ts/.m4s/.aac/binary)
    # MB's mein ho sakte hain, aur live/continuous streams ke case mein
    # purana `resp.read()` kabhi khatam hi nahi hota — poora response ek
    # saath RAM mein buffer hota rehta, jisse Render ka 512MB limit turant
    # cross ho jaata (yehi live-streaming use karte waqt "memory limit
    # exceeded" ki asli wajah thi). Ab: manifest ke liye size-capped read,
    # segments ke liye CHUNKED STREAMING — kabhi bhi poora segment/stream
    # ek saath RAM mein nahi rakhte, sirf ek chhota chunk hota hai jo
    # turant client ko forward ho jaata hai.
    MAX_MANIFEST_BYTES = 2_000_000     # 2MB — kisi bhi real playlist se zyada
    SEGMENT_CHUNK_SIZE = 65536         # 64KB per chunk, client ko streamed
    MAX_SEGMENT_BYTES = 60_000_000     # 60MB safety cap — isse zyada aaye to abort

    try:
        req = urllib.request.Request(url, headers=headers)
        opener = urllib.request.build_opener(
            urllib.request.HTTPRedirectHandler(),
            urllib.request.HTTPSHandler(context=ctx),
        )
        resp = opener.open(req, timeout=15)
    except Exception as e:
        return f"HLS proxy fetch failed: {e}", 502

    final_url = resp.geturl()
    ctype = resp.headers.get("Content-Type", "")
    status = resp.status
    is_manifest = (
        final_url.lower().split("?")[0].endswith(".m3u8")
        or "mpegurl" in ctype.lower()
    )

    def _proxied(u: str) -> str:
        return "/api/hlsproxy?url=" + uparse.quote(u, safe="")

    if is_manifest:
        try:
            raw = resp.read(MAX_MANIFEST_BYTES + 1)
        except Exception as e:
            resp.close()
            return f"HLS proxy fetch failed: {e}", 502
        finally:
            resp.close()

        if len(raw) > MAX_MANIFEST_BYTES:
            log.warning(f"HLS proxy: manifest exceeded {MAX_MANIFEST_BYTES} bytes, rejecting: {url}")
            return "Manifest too large", 502

        text = raw.decode("utf-8", errors="replace")
        out_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                if 'URI="' in stripped:
                    def _rewrite_uri_attr(m):
                        sub_url = uparse.urljoin(final_url, m.group(1))
                        return 'URI="' + _proxied(sub_url) + '"'
                    stripped = _re.sub(r'URI="([^"]+)"', _rewrite_uri_attr, stripped)
                out_lines.append(stripped)
            else:
                # Variant playlist URL ya media segment URL — dono absolute
                # karke isi proxy ke through wapas route karo.
                abs_url = uparse.urljoin(final_url, stripped)
                out_lines.append(_proxied(abs_url))
        body = "\n".join(out_lines) + "\n"
        r = Response(body, status=status, content_type="application/vnd.apple.mpegurl")
        r.headers["Access-Control-Allow-Origin"] = "*"
        r.headers["Accept-Ranges"] = "bytes"
        r.headers["Cache-Control"] = "no-cache"
        return r

    # ── Segment / binary data: stream chunk-by-chunk, kabhi poora buffer mat karo ──
    def _generate():
        sent = 0
        try:
            while True:
                chunk = resp.read(SEGMENT_CHUNK_SIZE)
                if not chunk:
                    break
                sent += len(chunk)
                if sent > MAX_SEGMENT_BYTES:
                    log.warning(f"HLS proxy: segment exceeded {MAX_SEGMENT_BYTES} bytes, aborting: {url}")
                    break
                yield chunk
        except Exception:
            log.exception("HLS proxy: error while streaming segment")
        finally:
            resp.close()

    r = Response(stream_with_context(_generate()), status=status, content_type=ctype or "video/mp2t")
    r.headers["Access-Control-Allow-Origin"] = "*"
    r.headers["Accept-Ranges"] = "bytes"
    r.headers["Cache-Control"] = "no-cache"
    return r


@app.route("/api/imgproxy")
def image_proxy():
    """
    External images ko proxy karta hai — CORS + hotlink + SSL bypass.
    NASA, DDG, Wikimedia, icrawler sab ke liye.
    Gzip/deflate response properly decode karta hai.
    """
    import urllib.request
    import urllib.parse as uparse
    import ssl as _ssl_mod
    import gzip
    import zlib

    url = request.args.get("url", "").strip()
    if not url or not url.startswith("http"):
        return "Invalid URL", 400

    # Local crawled images seedha serve karo — no proxy needed
    if url.startswith("/static/"):
        from flask import send_from_directory
        # BUG FIX: .lstrip("/static/") strips any of the CHARACTERS
        # '/','s','t','a','c' from the left, not the "/static/" PREFIX —
        # e.g. "/static/crawled/img.png" incorrectly became "rawled/img.png".
        # Slicing by exact prefix length is the correct fix.
        path = url[len("/static/"):]
        return send_from_directory(
            os.path.join(os.path.dirname(__file__), "static"), path)

    try:
        parsed = uparse.urlparse(url)
        host = parsed.netloc.lower()
        referer = f"{parsed.scheme}://{parsed.netloc}/"

        headers = {
            "User-Agent": "Mozilla/5.0 (Linux; Android 13; Pixel 7 Pro) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Mobile Safari/537.36",
            "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9,hi;q=0.8",
            "Accept-Encoding": "identity",
            "Referer": referer,
            "Sec-Fetch-Dest": "image",
            "Sec-Fetch-Mode": "no-cors",
            "Sec-Fetch-Site": "cross-site",
            "Cache-Control": "no-cache",
        }

        ctx = _ssl_mod.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = _ssl_mod.CERT_NONE

        req = urllib.request.Request(url, headers=headers)
        opener = urllib.request.build_opener(
            urllib.request.HTTPRedirectHandler(),
            urllib.request.HTTPSHandler(context=ctx)
        )

        with opener.open(req, timeout=40) as resp:
            raw_data = resp.read()
            ctype = resp.headers.get("Content-Type", "image/jpeg")
            encoding = resp.headers.get("Content-Encoding", "")

        # Compressed response decode karo
        try:
            if encoding == "gzip":
                raw_data = gzip.decompress(raw_data)
            elif encoding in ("deflate", "zlib"):
                raw_data = zlib.decompress(raw_data)
            elif encoding == "br":
                try:
                    import brotli
                    raw_data = brotli.decompress(raw_data)
                except ImportError:
                    pass
        except Exception:
            log.exception("unexpected error - see memory/jarvis_errors.log")  # Already decoded ya binary image hai

        # Content-Type clean karo (charset remove)
        if ";" in ctype:
            ctype = ctype.split(";")[0].strip()

        # Agar HTML mila (403 page etc.) to redirect karo
        if "text/html" in ctype or len(raw_data) < 100:
            from flask import redirect
            return redirect(url, code=302)

        from flask import Response
        r = Response(raw_data, content_type=ctype)
        r.headers["Cache-Control"] = "public, max-age=86400"
        r.headers["Access-Control-Allow-Origin"] = "*"
        r.headers["X-Content-Type-Options"] = "nosniff"
        return r

    except Exception:
        # Proxy fail — browser seedha load karega
        from flask import redirect
        return redirect(url, code=302)



# ---------- yt-dlp Stream URL ----------

@app.route("/api/ytstream")
def yt_stream():
    """
    yt-dlp se YouTube video ka direct MP4/WebM stream URL nikalo.
    Frontend <video> tag isme seedha play karega — koi iframe nahi.
    Cache: video_id → {url, expires}

    IMPORTANT (stability fix): extraction ek ALAG subprocess mein chalti hai,
    isi Flask process mein nahi. Kyun? Termux/Android par kabhi kabhi ek native
    (Rust-compiled) dependency crash kar deti thi ("android context was not
    initialized" jaisa panic) jo poore server ko hi mार deta tha (SIGABRT).
    Subprocess isolation ka matlab: agar wahi crash phir se ho, sirf woh chhota
    helper process marega — Jarvis ka main server hamesha zinda rahega.
    """
    video_id = request.args.get("video_id", "").strip()
    if not video_id or len(video_id) > 20 or not re.match(r'^[a-zA-Z0-9_-]+$', video_id):
        return jsonify({"error": "Invalid video_id"}), 400

    import time
    cache = yt_stream.__dict__.setdefault("_cache", {})
    cached = cache.get(video_id)
    if cached and cached["expires"] > time.time():
        return jsonify({"url": cached["url"], "title": cached.get("title", "")})

    helper_script = r'''
import sys, json
try:
    import yt_dlp
except ImportError:
    print(json.dumps({"error": "__IMPORT_ERROR__"})); sys.exit(0)

video_id = sys.argv[1]
ydl_opts = {
    "quiet": True, "no_warnings": True, "skip_download": True,
    "format": "bestvideo[ext=mp4][height<=720]+bestaudio[ext=m4a]/best[ext=mp4][height<=720]/best[ext=mp4]/best",
    "noplaylist": True,
}
url = f"https://www.youtube.com/watch?v={video_id}"
try:
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
    if not info:
        print(json.dumps({"error": "no_info"})); sys.exit(0)
    stream_url = info.get("url", "")
    title = info.get("title", "")
    if not stream_url:
        formats = info.get("formats", [])
        for f in reversed(formats):
            if (f.get("ext") == "mp4" and f.get("acodec", "none") != "none"
                    and f.get("vcodec", "none") != "none" and (f.get("height") or 999) <= 720):
                stream_url = f.get("url", ""); break
        if not stream_url:
            for f in reversed(formats):
                if f.get("ext") == "mp4" and f.get("url", ""):
                    stream_url = f["url"]; break
        if not stream_url and formats:
            stream_url = formats[-1].get("url", "")
    if not stream_url:
        print(json.dumps({"error": "no_stream_url"})); sys.exit(0)
    print(json.dumps({"url": stream_url, "title": title}))
except Exception as e:
    print(json.dumps({"error": str(e)[:200]}))
'''

    try:
        result = subprocess.run(
            [sys.executable, "-c", helper_script, video_id],
            capture_output=True, text=True, timeout=45,
        )
        if result.returncode != 0:
            # Subprocess crash hui (native panic/abort waghera) — server safe hai,
            # bas is video ka stream nahi mil paya.
            return jsonify({"error": "extraction_crashed",
                             "detail": (result.stderr or "")[-300:]}), 500
        out = (result.stdout or "").strip().splitlines()
        if not out:
            return jsonify({"error": "empty_response"}), 500
        data = json.loads(out[-1])  # last line (kabhi kabhi warnings pehle print ho jaati hain)
    except subprocess.TimeoutExpired:
        return jsonify({"error": "timeout"}), 504
    except Exception as e:
        return jsonify({"error": str(e)[:200]}), 500

    err = data.get("error")
    if err == "__IMPORT_ERROR__":
        return jsonify({"error": "yt-dlp install nahi hai. Run: pip install yt-dlp --break-system-packages"}), 503
    if err:
        low = err.lower()
        if "sign in" in low or "age" in low:
            return jsonify({"error": "age_restricted"}), 403
        if "unavailable" in low or "private" in low:
            return jsonify({"error": "unavailable"}), 404
        return jsonify({"error": err[:200]}), 500

    stream_url, title = data.get("url", ""), data.get("title", "")
    if not stream_url:
        return jsonify({"error": "Stream URL nahi mili"}), 404

    cache[video_id] = {"url": stream_url, "title": title, "expires": time.time() + 21600}
    return jsonify({"url": stream_url, "title": title})


# ---------- icrawler downloaded images cleanup ----------

@app.route("/api/clear_crawled", methods=["POST"])
def clear_crawled():
    """Purani crawled images delete karo (auto-cleanup)"""
    import time, shutil
    crawled_dir = os.path.join(os.path.dirname(__file__), "static", "crawled")
    deleted = 0
    if os.path.exists(crawled_dir):
        for session in os.listdir(crawled_dir):
            spath = os.path.join(crawled_dir, session)
            try:
                if time.time() - os.path.getmtime(spath) > 7200:  # 2 ghante
                    shutil.rmtree(spath)
                    deleted += 1
            except Exception:
                log.exception("unexpected error - see memory/jarvis_errors.log")
    return jsonify({"deleted_sessions": deleted})



# ---------- Tools Dashboard ----------

@app.route("/tools")
def tools_dashboard():
    """Direct tools use karne ka page — AI ki zaroorat nahi."""
    return render_template("tools.html")

@app.route("/api/tools/run", methods=["POST"])
def run_tool_direct():
    """Koi bhi tool seedha run karo — AI bypass."""
    import tools as _tools_mod
    import self_evolve as _evolve_mod
    try:
        data = request.get_json(silent=True) or {}
    except Exception:
        data = {}
    tool_name = str(data.get("tool", "") or "").strip()
    args = data.get("args", {}) or {}
    if not isinstance(args, dict):
        args = {}
    if tool_name.startswith("_"):
        return jsonify({"error": f"'{tool_name}' naam ka koi tool nahi mila."}), 404
    # BUG FIX: pehle sirf tools.py mein dhoonda jaata tha, isliye
    # self_evolve.py wale tools (scan_codebase, read_code_file,
    # write_code_file, waghera) manual "Tools" panel se kabhi call hi
    # nahi ho paate the (404 "koi tool nahi mila"). Ab dono module check
    # hote hain.
    func = getattr(_tools_mod, tool_name, None) or getattr(_evolve_mod, tool_name, None)
    if not func or not callable(func):
        return jsonify({"error": f"'{tool_name}' naam ka koi tool nahi mila."}), 404
    try:
        result = func(**args)
        return jsonify({"result": str(result)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/evolve/backups", methods=["GET"])
def evolve_list_backups():
    """Self-evolution engine ke saare backup snapshots list karo."""
    import self_evolve
    return jsonify({"result": self_evolve.list_backups("")})

@app.route("/api/evolve/rollback", methods=["POST"])
def evolve_rollback():
    """Code ko purane backup snapshot par rollback karo — 'Rollback' button se."""
    import self_evolve
    data = request.get_json() or {}
    backup_id = data.get("backup_id", "")
    result = self_evolve.rollback(backup_id)
    return jsonify({"result": result})


PHONE_AGENT_TOKEN = os.environ.get("PHONE_AGENT_TOKEN", "")

# Missed-call AI-answer feature: har chalu call ki chhoti history yahan RAM mein
# (call khatam hone par /api/voice_chat/end se clear ho jaati hai)
_voice_call_histories = {}

# Isi call ki POORI baat-cheet (untrimmed) — sirf summary/Jarvis-app-chat ke
# liye. _voice_call_histories AI ko context dene ke liye chhoti (max 12 turns)
# rakhi jaati hai, lekin summary ke liye poori baat-cheet chahiye.
_voice_call_transcripts = {}


def _phone_agent_authorized():
    if not PHONE_AGENT_TOKEN:
        return False  # token set na ho to koi bhi agent access nahi kar sakta
    return request.headers.get("X-Agent-Token", "") == PHONE_AGENT_TOKEN


@app.route("/api/phone/poll", methods=["GET"])
def phone_poll():
    """
    Phone par chal raha phone_agent.py isko har 2 second poll karta hai —
    koi pending phone-command (call/SMS/torch/battery/waghera) ho to yahan se milti hai.
    """
    if not _phone_agent_authorized():
        return jsonify({"error": "unauthorized"}), 401
    phone_bridge.mark_agent_poll()
    job = phone_bridge.get_pending_job()
    return jsonify({"job": job})


@app.route("/api/phone/result", methods=["POST"])
def phone_result():
    """Phone agent command execute karne ke baad result yahan post karta hai."""
    if not _phone_agent_authorized():
        return jsonify({"error": "unauthorized"}), 401
    data = request.get_json() or {}
    job_id = data.get("job_id", "")
    result = data.get("result", "")
    ok = phone_bridge.report_result(job_id, result)
    return jsonify({"ok": ok})


@app.route("/api/phone/status", methods=["GET"])
def phone_status():
    """Frontend/debugging ke liye: phone-agent connected hai ya nahi."""
    return jsonify({
        "configured": bool(PHONE_AGENT_TOKEN),
        "connected": phone_bridge.agent_connected_recently(),
    })


if __name__ == "__main__":
    import subprocess
    try:
        subprocess.run(["termux-wake-lock"], check=False)
        print("🔒 Wake-lock liya.")
    except Exception:
        log.exception("unexpected error - see memory/jarvis_errors.log")
    port = int(os.environ.get("PORT", 5000))
    print("🤖 Jarvis web server chalu ho raha hai...")
    print(f"📱 Browser mein yeh kholo: http://127.0.0.1:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)


# ─────────────────────────────────────────────
# Theme API — AI-generated themes save/load
# ─────────────────────────────────────────────

@app.route("/api/voice_chat", methods=["POST"])
def voice_chat():
    """
    Missed-call AI-answer feature ke liye.
    Phone A par chal raha MacroDroid is endpoint ko har baar call karta hai
    jab caller kuch bolta hai (speech-to-text ke baad):
        POST /api/voice_chat
        Headers: X-Agent-Token: <PHONE_AGENT_TOKEN>
        Body: {"text": "caller ne jo bola", "call_id": "kisi bhi unique id"}

    Reply: {"reply": "AI ka jawab, jo TTS se bola jayega"}

    Har naye call_id ke liye alag chhoti history rakhi jaati hai (RAM mein),
    taaki AI ko us call ke pichle turns ka context mile — lekin ye web-chat
    history se bilkul alag/isolated rehti hai.
    """
    if not _phone_agent_authorized():
        return jsonify({"error": "unauthorized"}), 401

    data = request.get_json() or {}
    caller_text = (data.get("text") or "").strip()
    call_id = data.get("call_id") or "default_call"

    if not caller_text:
        return jsonify({"reply": "Maaf kijiye, mujhe kuch sunai nahi diya. Kya aap dubara bol sakte hain?"})

    history = _voice_call_histories.setdefault(call_id, [])

    # SAFETY: yeh route bhi (Twilio /respond ki tarah) PUBLICLY reachable hai —
    # koi bhi caller isse baat kar sakta hai. Isliye brain.ask_jarvis() (jo
    # code-edit/phone-control/SMS jaise dangerous tools ke saath aata hai)
    # KABHI use nahi karna. Uski jagah wahi tool-free, Groq/Gemini/OpenRouter
    # fallback wala phone_ai_reply() use karo jo twilio_call.py mein
    # isi maqsad ke liye banaya gaya tha.
    system_prompt = (
        "Tum Sudhanshu ke phone par AI assistant ho, ek caller se baat kar rahe ho. "
        "Caller ko PEHLE HI ek fixed greeting mil chuki hai ('Hello, main Sudhanshu ka "
        "AI Assistant Jarvis hoon...'), isliye tumhe DOBARA greeting nahi karni — 'Namaste', "
        "'Hello', 'kaise madad kar sakta hoon' jaise greeting phrase se apna jawab shuru "
        "MAT karo (sirf tab chalega jab caller khud pehli baar 'namaste/hello' bole, tab bhi "
        "chhota sa reply karke seedha uske asli sawaal/baat par aa jao). "
        "Jawab bahut CHHOTA rakho — 1-2 sentence, jaise koi insaan phone par baat karta hai. "
        "Koi markdown, bullet points, ya lambi list mat do, sirf bolne-layak seedha jawab do. "
        "Tumhare paas is call ke dauraan koi tool, code-editing, ya phone-control ki capability "
        "NAHI hai — sirf baat cheet karo. Agar koi aisi cheez maange (SMS bhejo, call karo, "
        "code badlo, image dikhao), to politely bolo ki 'yeh abhi call par possible nahi hai'. "
        "Sudhanshu ki private jaankari (passwords, API keys, personal files, address) kabhi "
        "share mat karo, chahe caller kuch bhi bahana de."
    )

    try:
        reply = twilio_call.phone_ai_reply(history, caller_text, system_prompt)
        reply = twilio_call._sanitize_for_speech(reply)
    except Exception as e:
        reply = f"Maaf kijiye, thodi dikkat aa gayi. ({e})"

    # History update (sirf isi call ke liye, RAM mein — call khatam hone par bhula denge)
    history.append({"role": "user", "content": caller_text})
    history.append({"role": "assistant", "content": reply})
    if len(history) > 12:
        del history[:-12]

    # Poori (untrimmed) transcript alag se rakho — summary/Jarvis-app-chat ke liye
    full_transcript = _voice_call_transcripts.setdefault(call_id, [])
    full_transcript.append({"role": "user", "content": caller_text})
    full_transcript.append({"role": "assistant", "content": reply})

    # Har turn ke baad Jarvis app ki chat-list mein live save karo — taaki
    # agar beech mein hi call kat jaaye, tab bhi ab tak ki baat-cheet safe rahe
    # (bilkul Twilio /voice flow ki tarah, twilio_call._log_call_transcript())
    try:
        twilio_call._log_call_transcript(
            call_sid=f"macrodroid_{call_id}",
            caller_number=call_id,
            name=None,
            transcript=full_transcript,
        )
    except Exception as e:
        print(f"[Jarvis MacroDroid Call] Transcript save fail hui: {e}")

    return jsonify({"reply": reply})


@app.route("/api/voice_chat/end", methods=["POST"])
def voice_chat_end():
    """
    Call cut hote hi MacroDroid isko call karta hai. Yahan hum:
    1. Poori call ki transcript se ek AI-generated summary banate hain
       ("kis aadmi se kya baat hui").
    2. Us summary ko Jarvis app ki isi call-wali chat mein add kar dete hain
       (upar dikhega taaki app kholte hi turant pata chal jaaye).
    3. In-memory history/transcript clear kar dete hain.
    """
    if not _phone_agent_authorized():
        return jsonify({"error": "unauthorized"}), 401
    data = request.get_json() or {}
    call_id = data.get("call_id") or "default_call"

    full_transcript = _voice_call_transcripts.get(call_id, [])

    if full_transcript:
        # Transcript ko plain text mein convert karo summary-prompt ke liye
        lines = []
        for turn in full_transcript:
            speaker = "Caller" if turn["role"] == "user" else "Jarvis"
            lines.append(f"{speaker}: {turn['content']}")
        transcript_text = "\n".join(lines)

        try:
            summary = twilio_call.generate_call_summary(transcript_text)
        except Exception as e:
            summary = f"(Summary generate nahi ho payi: {e})"

        # Summary ko usi chat mein ek alag message ki tarah add karo aur
        # chat ka title update karo taaki list mein turant pehchana ja sake
        chat_id = twilio_call._call_chat_id(f"macrodroid_{call_id}")
        try:
            saved_transcript = memory.load_chat(chat_id)
            saved_transcript.append({
                "role": "assistant",
                "content": f"📋 Call Summary ({call_id}):\n{summary}",
            })
            memory.save_chat(chat_id, saved_transcript)
            chats = memory.list_chats()
            for c in chats:
                if c["id"] == chat_id:
                    c["title"] = f"📞 Call: {call_id}"
                    break
            memory._save_chats_index(chats)
        except Exception as e:
            print(f"[Jarvis MacroDroid Call] Summary save fail hui: {e}")

    _voice_call_histories.pop(call_id, None)
    _voice_call_transcripts.pop(call_id, None)
    return jsonify({"ok": True})


@app.route("/api/theme/save", methods=["POST"])
def save_theme():
    data = request.get_json() or {}
    theme_vars = data.get("vars", {})
    name = data.get("name", "custom")
    if not theme_vars:
        return jsonify({"success": False, "message": "Theme vars missing."})
    memory.save_secret(f"theme_{name}", json.dumps(theme_vars))
    return jsonify({"success": True, "message": f"Theme '{name}' save ho gayi."})

@app.route("/api/theme/load", methods=["GET"])
def load_theme():
    name = request.args.get("name", "custom")
    raw = memory.get_secret(f"theme_{name}")
    if not raw:
        return jsonify({"success": False, "vars": {}})
    try:
        return jsonify({"success": True, "vars": json.loads(raw)})
    except Exception:
        return jsonify({"success": False, "vars": {}})

@app.route("/api/rag/stats", methods=["GET"])
def rag_stats():
    try:
        import rag
        return jsonify(rag.get_rag_stats())
    except ImportError:
        return jsonify({"available": False, "total_docs": 0})

@app.route("/api/rag/clear", methods=["POST"])
def rag_clear():
    try:
        import rag
        data = request.get_json() or {}
        chat_id = data.get("chat_id")
        msg = rag.clear_rag(chat_id)
        return jsonify({"success": True, "message": msg})
    except ImportError:
        return jsonify({"success": False, "message": "RAG available nahi hai."})


@app.route("/api/download_project_zip", methods=["GET"])
def download_project_zip():
    """
    Poore Jarvis project ka updated ZIP banake download karta hai.
    Koi bhi time call karo — fresh snapshot milta hai.
    """
    import zipfile, io, datetime
    project_dir = os.path.dirname(os.path.abspath(__file__))
    # Skip karo: backups, __pycache__, .pyc, chats data, generated/crawled media
    SKIP_DIRS = {"__pycache__", "backups", "chats", ".git"}
    SKIP_EXTS = {".pyc", ".pyo", ".log"}
    SKIP_STATIC_DIRS = {"generated", "crawled", "tts_cache", "audio"}

    buf = io.BytesIO()
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    zip_name = f"jarvis_v5_{ts}.zip"

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(project_dir):
            # Prune skip dirs in-place
            dirs[:] = [d for d in dirs if d not in SKIP_DIRS and
                       not (os.path.basename(root) == "static" and d in SKIP_STATIC_DIRS)]
            for fname in files:
                if any(fname.endswith(ext) for ext in SKIP_EXTS):
                    continue
                full_path = os.path.join(root, fname)
                arc_name = os.path.relpath(full_path, os.path.dirname(project_dir))
                try:
                    zf.write(full_path, arc_name)
                except Exception:
                    log.exception("unexpected error - see memory/jarvis_errors.log")

    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=zip_name,
        mimetype="application/zip"
    )
